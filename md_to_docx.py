#!/usr/bin/env python3
"""
Markdown to DOCX Converter
A robust utility to convert Markdown files to DOCX format with advanced formatting support.

Author: Generated for church membership processing workflow
Date: 2025-06-10
"""

import os
import sys
import argparse
import logging
from pathlib import Path
from typing import List, Optional
import re
from datetime import datetime

try:
    import markdown
    from markdown.extensions import tables, codehilite, footnotes, toc
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError as e:
    print(f"Required package not found: {e}")
    print("Please install required packages:")
    print("pip install python-docx markdown")
    sys.exit(1)


class MarkdownToDocxConverter:
    """Convert Markdown files to DOCX with advanced formatting."""
    
    def __init__(self, verbose: bool = False, force_overwrite: bool = False):
        self.verbose = verbose
        self.force_overwrite = force_overwrite
        self.setup_logging()
        
    def setup_logging(self):
        """Setup logging configuration."""
        level = logging.DEBUG if self.verbose else logging.INFO
        logging.basicConfig(
            level=level,
            format='%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        self.logger = logging.getLogger(__name__)
    
    def get_markdown_files(self, directory: Path, specific_file: Optional[str] = None) -> List[Path]:
        """Get list of Markdown files to process."""
        if specific_file:
            file_path = Path(specific_file)
            if not file_path.exists():
                # Try relative to directory
                file_path = directory / specific_file
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {specific_file}")
            if file_path.suffix.lower() != '.md':
                raise ValueError(f"File is not a Markdown file: {specific_file}")
            return [file_path]
        
        # Get all .md files in directory
        md_files = list(directory.glob('*.md'))
        if not md_files:
            raise FileNotFoundError(f"No Markdown files found in {directory}")
        
        return sorted(md_files)
    
    def should_skip_file(self, md_file: Path, output_dir: Path) -> bool:
        """Check if file should be skipped (already converted and not forcing)."""
        if self.force_overwrite:
            return False
            
        docx_file = output_dir / f"{md_file.stem}.docx"
        if docx_file.exists():
            self.logger.info(f"Skipping {md_file.name} - DOCX already exists")
            return True
        return False
    
    def parse_markdown(self, content: str) -> str:
        """Parse Markdown content to HTML with extensions."""
        md = markdown.Markdown(extensions=[
            'tables',
            'codehilite',
            'footnotes',
            'toc',
            'fenced_code',
            'attr_list',
            'def_list',
            'abbr',
            'md_in_html'
        ])
        return md.convert(content)
    
    def setup_document_styles(self, doc: Document):
        """Setup custom styles for the document."""
        styles = doc.styles
        
        # Code block style
        try:
            code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
            code_font = code_style.font
            code_font.name = 'Consolas'
            code_font.size = Pt(9)
            code_style.paragraph_format.left_indent = Inches(0.5)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
        except ValueError:
            # Style already exists
            pass
        
        # Quote style
        try:
            quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_style.paragraph_format.left_indent = Inches(0.5)
            quote_style.paragraph_format.right_indent = Inches(0.5)
            quote_font = quote_style.font
            quote_font.italic = True
        except ValueError:
            # Style already exists
            pass
    
    def html_to_docx(self, html_content: str, doc: Document):
        """Convert HTML content to DOCX format."""
        from html.parser import HTMLParser
        from docx.shared import Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        class DocxHTMLParser(HTMLParser):
            def __init__(self, document):
                super().__init__()
                self.doc = document
                self.current_paragraph = None
                self.current_run = None
                self.tag_stack = []
                self.list_level = 0
                self.in_code_block = False
                self.code_content = []
                
                # Table handling
                self.current_table = None
                self.current_row = None
                self.current_cell = None
                self.table_data = []
                self.current_row_data = []
                self.current_cell_content = []
                self.in_table = False
                self.in_table_header = False
                
            def handle_starttag(self, tag, attrs):
                self.tag_stack.append(tag)
                
                if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    self.current_paragraph = self.doc.add_heading(level=int(tag[1]))
                elif tag == 'p':
                    if not self.in_table:
                        self.current_paragraph = self.doc.add_paragraph()
                elif tag == 'pre':
                    self.in_code_block = True
                    self.code_content = []
                elif tag == 'code' and not self.in_code_block:
                    if not self.in_table:
                        if self.current_paragraph is None:
                            self.current_paragraph = self.doc.add_paragraph()
                        self.current_run = self.current_paragraph.add_run()
                        self.current_run.font.name = 'Consolas'
                elif tag == 'strong' or tag == 'b':
                    if not self.in_table:
                        if self.current_paragraph is None:
                            self.current_paragraph = self.doc.add_paragraph()
                        self.current_run = self.current_paragraph.add_run()
                        self.current_run.bold = True
                elif tag == 'em' or tag == 'i':
                    if not self.in_table:
                        if self.current_paragraph is None:
                            self.current_paragraph = self.doc.add_paragraph()
                        self.current_run = self.current_paragraph.add_run()
                        self.current_run.italic = True
                elif tag == 'ul' or tag == 'ol':
                    self.list_level += 1
                elif tag == 'li':
                    if not self.in_table:
                        self.current_paragraph = self.doc.add_paragraph(style='List Bullet' if self.tag_stack.count('ul') > 0 else 'List Number')
                elif tag == 'blockquote':
                    if not self.in_table:
                        self.current_paragraph = self.doc.add_paragraph(style='Quote')
                elif tag == 'table':
                    self.in_table = True
                    self.table_data = []
                elif tag == 'thead':
                    self.in_table_header = True
                elif tag == 'tbody':
                    self.in_table_header = False
                elif tag == 'tr':
                    self.current_row_data = []
                elif tag == 'th' or tag == 'td':
                    self.current_cell_content = []
                elif tag == 'br':
                    if self.in_table:
                        self.current_cell_content.append('\n')
                    elif self.current_paragraph:
                        self.current_paragraph.add_run().add_break()
            
            def handle_endtag(self, tag):
                if self.tag_stack and self.tag_stack[-1] == tag:
                    self.tag_stack.pop()
                
                if tag == 'pre':
                    if self.code_content:
                        code_para = self.doc.add_paragraph('\n'.join(self.code_content))
                        code_para.style = 'Code Block'
                    self.in_code_block = False
                    self.code_content = []
                elif tag == 'table':
                    self._create_docx_table()
                    self.in_table = False
                elif tag == 'thead':
                    self.in_table_header = False
                elif tag == 'tr':
                    if self.current_row_data:
                        self.table_data.append({
                            'data': self.current_row_data.copy(),
                            'is_header': self.in_table_header
                        })
                    self.current_row_data = []
                elif tag == 'th' or tag == 'td':
                    cell_text = ''.join(self.current_cell_content).strip()
                    self.current_row_data.append(cell_text)
                    self.current_cell_content = []
                elif tag in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'blockquote']:
                    if not self.in_table:
                        self.current_paragraph = None
                        self.current_run = None
                elif tag in ['strong', 'b', 'em', 'i', 'code']:
                    if not self.in_table:
                        self.current_run = None
                elif tag in ['ul', 'ol']:
                    self.list_level = max(0, self.list_level - 1)
            
            def handle_data(self, data):
                if self.in_code_block:
                    self.code_content.append(data)
                elif self.in_table:
                    self.current_cell_content.append(data)
                elif data.strip():
                    if self.current_paragraph is None:
                        self.current_paragraph = self.doc.add_paragraph()
                    
                    if self.current_run is None:
                        self.current_run = self.current_paragraph.add_run()
                    
                    self.current_run.text += data
            
            def _calculate_column_widths(self, table_data):
                """Calculate optimal column widths based on content length."""
                if not table_data:
                    return []
                
                max_cols = max(len(row['data']) for row in table_data)
                col_max_lengths = [0] * max_cols
                
                # Calculate maximum content length per column
                for row in table_data:
                    for i, cell in enumerate(row['data']):
                        if i < max_cols:
                            # Consider line breaks in length calculation
                            lines = cell.split('\n')
                            max_line_length = max(len(line) for line in lines) if lines else 0
                            col_max_lengths[i] = max(col_max_lengths[i], max_line_length)
                
                # Calculate proportional widths
                total_length = sum(col_max_lengths)
                if total_length == 0:
                    return [Inches(6.0 / max_cols)] * max_cols
                
                # Available width (assuming 6 inches total table width)
                available_width = 6.0
                min_col_width = 0.8  # Minimum column width in inches
                
                # Calculate proportional widths
                widths = []
                for length in col_max_lengths:
                    if total_length > 0:
                        proportional_width = (length / total_length) * available_width
                        # Ensure minimum width
                        width = max(proportional_width, min_col_width)
                        widths.append(Inches(width))
                    else:
                        widths.append(Inches(min_col_width))
                
                # Adjust if total width exceeds available space
                total_width = sum(w.inches for w in widths)
                if total_width > available_width:
                    scale_factor = available_width / total_width
                    widths = [Inches(w.inches * scale_factor) for w in widths]
                
                return widths
            
            def _create_docx_table(self):
                """Create a DOCX table from collected table data."""
                if not self.table_data:
                    return
                
                # Determine table dimensions
                max_cols = max(len(row['data']) for row in self.table_data)
                num_rows = len(self.table_data)
                
                # Create table
                table = self.doc.add_table(rows=num_rows, cols=max_cols)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                
                # Calculate and set column widths
                col_widths = self._calculate_column_widths(self.table_data)
                for i, width in enumerate(col_widths):
                    if i < len(table.columns):
                        table.columns[i].width = width
                
                # Populate table
                for row_idx, row_data in enumerate(self.table_data):
                    table_row = table.rows[row_idx]
                    
                    for col_idx, cell_data in enumerate(row_data['data']):
                        if col_idx < max_cols:
                            cell = table_row.cells[col_idx]
                            
                            # Clear default paragraph and add content
                            cell.text = cell_data
                            
                            # Format header cells
                            if row_data['is_header']:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                                
                                # Add background color to header
                                self._set_cell_background(cell, "D9D9D9")
                            
                            # Enable text wrapping
                            self._enable_cell_text_wrap(cell)
                
                # Add table style
                table.style = 'Table Grid'
                
                # Add some spacing after table
                self.doc.add_paragraph()
            
            def _set_cell_background(self, cell, color):
                """Set background color for a table cell."""
                try:
                    from docx.oxml import parse_xml
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
                        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"', 
                        color))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                except:
                    # Fallback if shading fails
                    pass
            
            def _enable_cell_text_wrap(self, cell):
                """Enable text wrapping for table cell."""
                try:
                    # Set cell properties for text wrapping
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    # Enable text wrapping
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:type'), 'auto')
                    tcPr.append(tcW)
                    
                    # Set vertical alignment
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), 'top')
                    tcPr.append(vAlign)
                except:
                    # Fallback if XML manipulation fails
                    pass
        
        parser = DocxHTMLParser(doc)
        parser.feed(html_content)
    
    def convert_file(self, md_file: Path, output_dir: Path) -> bool:
        """Convert a single Markdown file to DOCX."""
        try:
            self.logger.info(f"Converting {md_file.name}...")
            
            # Read Markdown content
            with open(md_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Parse Markdown to HTML
            html_content = self.parse_markdown(md_content)
            
            # Create DOCX document
            doc = Document()
            self.setup_document_styles(doc)
            
            # Add title from filename
            title = md_file.stem.replace('_', ' ').replace('-', ' ').title()
            doc.add_heading(title, 0)
            
            # Convert HTML to DOCX
            self.html_to_docx(html_content, doc)
            
            # Save DOCX file
            output_file = output_dir / f"{md_file.stem}.docx"
            doc.save(output_file)
            
            self.logger.info(f"✓ Converted to {output_file.name}")
            return True
            
        except Exception as e:
            self.logger.error(f"✗ Failed to convert {md_file.name}: {str(e)}")
            return False
    
    def convert_files(self, input_path: str, output_dir: Optional[str] = None, 
                     specific_file: Optional[str] = None) -> tuple[int, int]:
        """Convert Markdown files to DOCX format."""
        try:
            # Setup paths
            input_dir = Path(input_path).resolve()
            if not input_dir.exists():
                raise FileNotFoundError(f"Input directory not found: {input_path}")
            
            output_directory = Path(output_dir) if output_dir else input_dir
            output_directory.mkdir(parents=True, exist_ok=True)
            
            # Get files to process
            md_files = self.get_markdown_files(input_dir, specific_file)
            
            self.logger.info(f"Found {len(md_files)} Markdown file(s) to process")
            
            # Process files
            converted = 0
            skipped = 0
            
            for md_file in md_files:
                if self.should_skip_file(md_file, output_directory):
                    skipped += 1
                    continue
                
                if self.convert_file(md_file, output_directory):
                    converted += 1
            
            # Summary
            total = len(md_files)
            failed = total - converted - skipped
            
            self.logger.info(f"\nConversion Summary:")
            self.logger.info(f"  Total files: {total}")
            self.logger.info(f"  Converted: {converted}")
            self.logger.info(f"  Skipped: {skipped}")
            if failed > 0:
                self.logger.info(f"  Failed: {failed}")
            
            return converted, failed
            
        except Exception as e:
            self.logger.error(f"Conversion failed: {str(e)}")
            return 0, 1


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description="Convert Markdown files to DOCX format with advanced formatting support",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s                          # Convert all .md files in current directory
  %(prog)s -f document.md           # Convert specific file
  %(prog)s -o /path/to/output       # Specify output directory
  %(prog)s -v --force               # Verbose mode with overwrite
  %(prog)s -d /path/to/markdown     # Process files in specific directory
        """
    )
    
    parser.add_argument(
        '-f', '--file',
        help='Specific Markdown file to convert'
    )
    
    parser.add_argument(
        '-d', '--directory',
        default='.',
        help='Input directory containing Markdown files (default: current directory)'
    )
    
    parser.add_argument(
        '-o', '--output',
        help='Output directory for DOCX files (default: same as input directory)'
    )
    
    parser.add_argument(
        '-v', '--verbose',
        action='store_true',
        help='Enable verbose output'
    )
    
    parser.add_argument(
        '--force',
        action='store_true',
        help='Force overwrite existing DOCX files'
    )
    
    parser.add_argument(
        '--version',
        action='version',
        version='Markdown to DOCX Converter v1.0'
    )
    
    args = parser.parse_args()
    
    # Create converter
    converter = MarkdownToDocxConverter(
        verbose=args.verbose,
        force_overwrite=args.force
    )
    
    # Convert files
    try:
        converted, failed = converter.convert_files(
            input_path=args.directory,
            output_dir=args.output,
            specific_file=args.file
        )
        
        if failed > 0:
            sys.exit(1)
        elif converted == 0:
            print("No files were converted.")
            sys.exit(1)
        else:
            print(f"Successfully converted {converted} file(s).")
            
    except KeyboardInterrupt:
        print("\nConversion interrupted by user.")
        sys.exit(1)
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)


if __name__ == '__main__':
    main()
