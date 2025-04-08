import os
import email
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from collections import defaultdict
from datetime import datetime
import pytz
import re
import sys

def format_date(date_str):
    # Parse the email date
    dt = email.utils.parsedate_to_datetime(date_str)
    
    # Convert to Eastern Time
    eastern = pytz.timezone('US/Eastern')
    dt_eastern = dt.astimezone(eastern)
    
    # Format with timezone abbreviation
    return dt_eastern.strftime('%B %d, %Y %I:%M %p %Z')

def extract_content(eml_file):
    with open(eml_file, 'r', encoding='utf-8') as f:
        msg = email.message_from_file(f)
    
    date = format_date(msg.get('Date', ''))
    
    for part in msg.walk():
        if part.get_content_type() == "text/html":
            return date, part.get_payload(decode=True).decode()
    return date, ""

def parse_html_content(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    table = soup.find('table', class_='content-table')
    if not table:
        return None, {}
    
    name_cell = table.find('td')
    name = name_cell.find('a').text.strip() if name_cell and name_cell.find('a') else "Unknown"
    
    questions = {
        "Email address": "",
        "Phone number": "",
        "Birthday": "",
        "Address": "",
        "Primary campus": "",
        "Did you regularly attend another church previously?": "",
        "Former Church Name (if applicable)": "",
        "Former Church Address": "",
        "Have you been baptized by immersion following salvation?": "",
        "Have you completed the Getting Connected/Membership class?": "",
        "Please type your full name here if you have read and agree with the church covenant, bylaws and statement of faith.": "",
        "Have you accepted Jesus Christ as Lord and Savior? Who is Jesus to you?": "",
        "When were you saved?": "",
        "Please describe the circumstances and events the led you to accept Christ as your Savior and Lord.": "",
        "Please describe your own concept of the meaning of water baptism.": "",
        "Your short bio": ""
    }
    
    for row in table.find_all('tr'):
        cells = row.find_all('td')
        if not cells:
            continue
            
        cell = cells[0]
        strong = cell.find('strong')
        if not strong:
            continue
            
        question = strong.text.strip()
        if question in questions:
            for elem in cell.find_all(['strong', 'small', 'div']):
                elem.decompose()
            # Handle line breaks
            for br in cell.find_all('br'):
                br.replace_with('\n')
            answer = cell.get_text()
            # Clean up but preserve intentional line breaks
            answer = '\n'.join(line.strip() for line in answer.split('\n'))
            # Remove empty lines while preserving paragraph breaks
            answer = '\n'.join(filter(None, answer.split('\n')))
            questions[question] = answer
    
    return name, questions

def normalize_name(name):
    """
    Normalize a name for comparison purposes.
    Remove comments, extra spaces, and convert to lowercase.
    """
    if "#" in name:
        name = name.split("#", 1)[0]
    return " ".join(name.lower().split())

def split_name_and_comment(name_line):
    """
    Split a line into name and comment parts.
    """
    if "#" in name_line:
        name, comment = name_line.split("#", 1)
        return name.strip(), "#" + comment
    return name_line.strip(), ""

def format_name_for_file(name, comments_dict=None):
    """
    Format a name for member_names.txt file.
    Handle Jr/Sr suffixes and preserve comments.
    """
    # Check for comments in the name
    comment = ""
    name_part = name
    if "#" in name:
        name_part, comment = name.split("#", 1)
        name_part = name_part.strip()
        comment = "#" + comment
    
    # Check for existing comment in the dictionary
    normalized = normalize_name(name_part)
    if comments_dict and normalized in comments_dict and not comment:
        comment = comments_dict[normalized]
    
    # Split name into parts
    parts = name_part.split()
    
    # Check for Jr/Sr suffixes
    suffix_match = re.search(r'\b(Jr|Sr|Jr\.|Sr\.)$', name_part, re.IGNORECASE)
    if suffix_match and len(parts) >= 3:
        # Remove suffix (could be at the end)
        suffix = parts.pop()
        # Remove period if present
        suffix = suffix.rstrip('.')
        # Get last name
        last_name = parts.pop()
        # Get first and middle names
        first_parts = " ".join(parts)
        # Reformat as "First Middle Jr Last"
        formatted_name = f"{first_parts} {suffix} {last_name}"
    else:
        formatted_name = name_part
    
    # Reattach comment if it exists
    if comment:
        formatted_name = f"{formatted_name} {comment}"
    
    return formatted_name

def get_sort_key(name):
    """
    Get sort key for a name, handling hyphenated last names
    and extracting the name part if there's a comment.
    """
    # Remove comment if exists
    if "#" in name:
        name = name.split("#", 1)[0].strip()
    
    # Split name into parts
    parts = name.split()
    
    # No parts (empty string)
    if not parts:
        return ""
    
    # Handle Jr/Sr suffixes - if the second part is Jr/Sr, the last name is the last part
    if len(parts) >= 3 and re.search(r'\b(Jr|Sr|Jr\.|Sr\.)\b', parts[1], re.IGNORECASE):
        return parts[-1]
    
    # Get the last part as the last name
    last_name = parts[-1]
    
    # Handle hyphenated last names - sort by second part
    if "-" in last_name:
        last_parts = last_name.split("-")
        return last_parts[-1]
    
    return last_name

def read_existing_file_data(names_file):
    """
    Read existing names file and extract names, comments, and submission line.
    Returns:
    - campus_names: Dictionary mapping campus to list of names
    - comments_dict: Dictionary mapping normalized names to comments
    - submission_line: The "Above Submitted by" line if it exists
    """
    campus_names = defaultdict(list)
    comments_dict = {}
    submission_line = None
    
    if not os.path.exists(names_file):
        return campus_names, comments_dict, submission_line
    
    with open(names_file, 'r') as f:
        current_campus = None
        for line in f:
            line = line.strip()
            if not line:
                continue
                
            # Detect campus headers
            if line.startswith("==") and line.endswith("=="):
                current_campus = line[2:-2].strip()
            # Detect submission line
            elif line.startswith("Above Submitted by"):
                submission_line = line
            # Process name entries
            elif current_campus:
                campus_names[current_campus].append(line)
                
                # Store comments for future reference
                name_part, comment_part = split_name_and_comment(line)
                if comment_part:
                    comments_dict[normalize_name(name_part)] = comment_part
    
    return campus_names, comments_dict, submission_line

def collect_member_data(eml_files):
    """
    Collect member data from all .eml files.
    Returns a list of (name, campus) tuples.
    """
    members_data = []
    
    for eml_file in eml_files:
        date, html_content = extract_content(eml_file)
        name, questions = parse_html_content(html_content)
        
        if name and "Primary campus" in questions:
            members_data.append((name, questions["Primary campus"]))
            
    return members_data

def update_names_file(members_data, replace_file=True):
    """
    Update member_names.txt with new names, either by replacing the file
    or appending to it depending on user's choice.
    """
    names_file = "member_names.txt"
    file_exists = os.path.exists(names_file)
    
    # Read existing file to capture names, comments, and submission line
    existing_campus_names, comments_dict, submission_line = read_existing_file_data(names_file)
    
    # Initialize campus_names based on mode
    if replace_file:
        # In replace mode, start with empty campus lists
        campus_names = defaultdict(list)
    else:
        # In append mode, use existing campus names
        campus_names = existing_campus_names
    
    # Process and add all member data
    for name, campus in members_data:
        # Format the name for the file, preserving comments
        formatted_name = format_name_for_file(name, comments_dict)
        
        # Add formatted name to appropriate campus
        if not replace_file:
            # Only add if not already in the list (for append mode)
            normalized = normalize_name(split_name_and_comment(formatted_name)[0])
            existing_names = [normalize_name(split_name_and_comment(n)[0]) for n in campus_names[campus]]
            if normalized not in existing_names:
                campus_names[campus].append(formatted_name)
        else:
            # Always add in replace mode
            campus_names[campus].append(formatted_name)
    
    # Write names back to file
    with open(names_file, 'w') as f:
        for campus_name in sorted(campus_names.keys()):
            f.write(f"=={campus_name}==\n")
            
            # Sort names within this campus
            if replace_file:
                # If replacing, sort all names
                sorted_names = sorted(campus_names[campus_name], key=get_sort_key)
            else:
                # If appending, maintain existing names in their original order
                sorted_names = campus_names[campus_name]
            
            for name_entry in sorted_names:
                f.write(name_entry + '\n')
            
            f.write('\n')
        
        # Add submission line
        today = datetime.now().strftime("%Y-%m-%d")
        if submission_line and not replace_file:
            f.write(submission_line + '\n')
        f.write(f"Above Submitted by {today}\n")

def create_document_files(eml_file, create_word=False):
    """
    Create PDF files for a member, and optionally DOCX files if create_word is True.
    """
    if not os.path.exists('WordDocs'):
        os.makedirs('WordDocs')
        
    date, html_content = extract_content(eml_file)
    name, questions = parse_html_content(html_content)
    
    name_parts = name.split()
    if len(name_parts) >= 2:
        lastname, firstname = name_parts[-1], name_parts[0]
    else: 
        lastname = firstname = "Unknown"
    
    base_filename = f"{lastname}_{firstname}"
    docx_filename = f"WordDocs/{base_filename}.docx"
    pdf_filename = f"WordDocs/{base_filename}.pdf"
    
    # Define is_short_answer function outside of conditional blocks
    def is_short_answer(answer):
        return len(answer) < 50 and '\n' not in answer
    
    # Create Word document if requested
    if create_word:
        doc = Document()
        doc.styles['Normal'].paragraph_format.line_spacing = 1.0
        doc.styles['Normal'].paragraph_format.space_after = Pt(0)
        doc.styles['Normal'].paragraph_format.space_before = Pt(0)
            
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(name)
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        # Add submission date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Submitted: {date}")
        date_run.font.size = Pt(10)
                
        # Add questions and answers
        for question, answer in questions.items():
            doc.add_paragraph()
            
            answer = answer.strip() if answer else "Not Provided"
            if question in ["Have you been baptized by immersion following salvation?",
                           "Have you completed the Getting Connected/Membership class?"]:
                answer = "Yes" if "yes" in answer.lower() else "No" if "no" in answer.lower() else answer
            
            if question in ["Email address", "Phone number", "Birthday", "Address", "Primary campus"]:
                p = doc.add_paragraph()
                q_run = p.add_run(question)
                q_run.bold = True
                p.add_run(f": {answer}")
                continue
                    
            p = doc.add_paragraph()
            q_run = p.add_run(question)
            q_run.bold = True
            
            if is_short_answer(answer):
                p.add_run(": " + answer)
            else:
                doc.add_paragraph(answer)
                    
        # Save Word document
        doc.save(docx_filename)
            
    # Create PDF
    doc = SimpleDocTemplate(pdf_filename, pagesize=letter,
                         leftMargin=36, rightMargin=36,
                         topMargin=36, bottomMargin=36)  # 36 points = 0.5 inches
    styles = getSampleStyleSheet()
    story = []
                 
    # Title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        fontSize=14,
        spaceAfter=20
    )
                 
    # Add title
    story.append(Paragraph(name, title_style))
    
    # Add submission date
    date_style = ParagraphStyle(
        'Date',
        parent=styles['Normal'],
        alignment=TA_CENTER,
        fontSize=10,
        spaceBefore=6,
        spaceAfter=12
    )
    story.append(Paragraph(f"Submitted: {date}", date_style))
            
    # Question style
    question_style = ParagraphStyle(
        'Question',
        parent=styles['Normal'],
        fontSize=11,
        fontName='Helvetica-Bold',
        spaceAfter=0,
        spaceBefore=12
    )
            
    # Answer style
    answer_style = ParagraphStyle(
        'Answer',
        parent=styles['Normal'],
        spaceAfter=0
    )
            
    # Footer style
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        alignment=TA_RIGHT,
        fontSize=8
    )
            
    for question, answer in questions.items():
        answer = answer.strip() if answer else "Not Provided"
        if question in ["Have you been baptized by immersion following salvation?",
                       "Have you completed the Getting Connected/Membership class?"]:
            answer = "Yes" if "yes" in answer.lower() else "No" if "no" in answer.lower() else answer
            
        if question in ["Email address", "Phone number", "Birthday", "Address", "Primary campus"]:
            story.append(Paragraph(f"<b>{question}</b>: {answer}", answer_style))
            continue
            
        if is_short_answer(answer):
            story.append(Paragraph(f"<b>{question}</b>: {answer}", answer_style))
        else:
            story.append(Paragraph(f"<b>{question}</b>", question_style))
            # Replace newlines with <br/> tags for PDF
            formatted_answer = answer.replace('\n', '<br/>')
            story.append(Paragraph(formatted_answer, answer_style))
             
    # Build PDF with footer on pages after first
    def add_footer(canvas, doc):
        if doc.page > 1:
            canvas.saveState()
            footer = Paragraph(f"{name} - Continued", footer_style)
            w, h = footer.wrap(doc.width, doc.bottomMargin)
            footer.drawOn(canvas, doc.leftMargin, h)
            canvas.restoreState()
         
    doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
    
    if create_word:
        return docx_filename, pdf_filename
    else:
        return None, pdf_filename

def process_eml_files(create_word=False):
    # Collect all .eml files
    eml_files = [filename for filename in os.listdir('.') if filename.endswith('.eml')]
    
    if not eml_files:
        print("No .eml files found in the current directory.")
        return
    
    # Prompt user once for file handling option
    names_file = "member_names.txt"
    replace_file = True
    
    if os.path.exists(names_file):
        while True:
            choice = input(f"\n{names_file} already exists. Replace (R) or Append (A)? ").upper()
            if choice in ['R', 'A']:
                replace_file = (choice == 'R')
                break
            print("Invalid choice. Please enter 'R' or 'A'.")
    
    # Collect member data from all files
    members_data = collect_member_data(eml_files)
    
    # Update the names file with all collected data
    update_names_file(members_data, replace_file)
    
    # Create document files for each .eml file
    for eml_file in eml_files:
        docx_file, pdf_file = create_document_files(eml_file, create_word)
        if create_word:
            print(f"Processed {eml_file} -> {docx_file}, {pdf_file}")
        else:
            print(f"Processed {eml_file} -> {pdf_file}")

if __name__ == "__main__":
    # Check for command line arguments
    create_word = "-w" in sys.argv
    process_eml_files(create_word)
